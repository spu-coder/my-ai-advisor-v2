import os
import logging
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
from typing import List, Dict, Any
from langchain_core.documents import Document

logger = logging.getLogger("DATA_PROCESSOR")

# ------------------------------------------------------------
# وظائف استخراج النص
# ------------------------------------------------------------

def _extract_text_from_pdf(file_path: str) -> str:
    """يستخرج النص من ملف PDF، بما في ذلك محاولة OCR للصور المضمنة."""
    full_text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # استخراج النص العادي
                text = page.extract_text()
                if text:
                    full_text += text + "\n\n"
                
                # محاولة OCR للصور المضمنة (يتطلب مكتبة خارجية مثل pdf2image و pytesseract)
                # نظراً لقيود البيئة، سنعتمد على النص المستخرج مباشرة من pdfplumber
                # يمكن إضافة دعم OCR المتقدم هنا باستخدام مكتبات مثل pdf2image و pytesseract
                
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
    return full_text

def _extract_text_from_docx(file_path: str) -> str:
    """يستخرج النص من ملف DOCX."""
    try:
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        return ""

def _extract_text_from_image(file_path: str) -> str:
    """يستخرج النص من ملف صورة باستخدام Tesseract OCR."""
    try:
        text = pytesseract.image_to_string(Image.open(file_path), lang='ara+eng')
        return text
    except pytesseract.TesseractNotFoundError:
        logger.error("Tesseract is not installed or not in your PATH. Cannot perform OCR.")
        return ""
    except Exception as e:
        logger.error(f"Error extracting text from image {file_path}: {e}")
        return ""

def _extract_text_from_txt(file_path: str) -> str:
    """يستخرج النص من ملف TXT."""
    try:
        # محاولة قراءة بترميزات مختلفة
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1256', 'windows-1256']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        # إذا فشلت جميع الترميزات، استخدم utf-8 مع errors='ignore'
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {e}")
        return ""

def process_document(file_path: str) -> Document | None:
    """
    يعالج ملفاً واحداً ويستخرج منه النص لإنشاء كائن Document.
    يدعم PDF، DOCX، DOC، TXT، والصور (JPG, PNG, TIFF).
    """
    filename = os.path.basename(file_path)
    full_text = ""
    
    if filename.lower().endswith((".pdf")):
        full_text = _extract_text_from_pdf(file_path)
    elif filename.lower().endswith((".docx", ".doc")):
        full_text = _extract_text_from_docx(file_path)
    elif filename.lower().endswith((".txt")):
        full_text = _extract_text_from_txt(file_path)
    elif filename.lower().endswith((".jpg", ".jpeg", ".png", ".tiff")):
        full_text = _extract_text_from_image(file_path)
    else:
        logger.warning(f"Unsupported file type for ingestion: {filename}")
        return None
        
    if full_text and full_text.strip():
        # إضافة معالجة الجداول المتقدمة هنا إذا لزم الأمر (يتطلب مكتبات إضافية مثل camelot أو tabula)
        # حالياً، يتم الاعتماد على النص المستخرج، والذي يجب أن يتضمن الجداول كنص عادي.
        return Document(page_content=full_text, metadata={"source": filename})
    
    return None

def ingest_all_documents(data_dir: str) -> List[Document]:
    """يفهرس جميع المستندات المدعومة في مجلد البيانات."""
    loaded_docs = []
    logger.info(f"Scanning directory: {data_dir}")
    
    if not os.path.exists(data_dir):
        logger.error(f"Data directory does not exist: {data_dir}")
        return loaded_docs
    
    files = os.listdir(data_dir)
    logger.info(f"Found {len(files)} files in directory")
    
    for filename in files:
        file_path = os.path.join(data_dir, filename)
        if os.path.isfile(file_path):
            logger.info(f"Processing file: {filename}")
            try:
                doc = process_document(file_path)
                if doc:
                    # التحقق من أن النص غير فارغ
                    if doc.page_content and doc.page_content.strip():
                        logger.info(f"Successfully processed {filename} - Content length: {len(doc.page_content)} characters")
                        loaded_docs.append(doc)
                    else:
                        logger.warning(f"File {filename} produced empty content")
                else:
                    logger.warning(f"Failed to process {filename}")
            except Exception as e:
                logger.error(f"Error processing {filename}: {e}", exc_info=True)
    
    logger.info(f"Total documents loaded: {len(loaded_docs)}")
    return loaded_docs
